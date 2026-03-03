#!/usr/bin/env python3
from __future__ import annotations

import re
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt


def _collapse_ws(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _extract_abstract(first_pages_text: str) -> str | None:
    text = _collapse_ws(first_pages_text)
    m = re.search(r"\bAbstract\b|\bABSTRACT\b", text, re.I)
    if not m:
        return None
    tail = text[m.end() :]
    end = None
    for pat in [
        r"\bKeywords\b",
        r"\bIndex Terms\b",
        r"\b1\.?\s+Introduction\b",
        r"\bI\.?\s+INTRODUCTION\b",
        r"\bINTRODUCTION\b",
        r"\bCategories and Subject Descriptors\b",
    ]:
        m2 = re.search(pat, tail, re.I)
        if m2:
            end = m2.start() if end is None else min(end, m2.start())
    return tail[:end].strip() if end is not None else tail.strip()


def _extract_conclusion(full_text: str) -> str | None:
    text = _collapse_ws(full_text)
    # Prefer numbered section headings, then generic "Conclusions"
    matches = list(re.finditer(r"\b\d+\s*\.?\s+Conclusions?\b", text, re.I))
    if not matches:
        matches = list(re.finditer(r"\bConclusions?\b", text, re.I))
    if not matches:
        return None
    m = matches[-1]
    tail = text[m.end() : m.end() + 50000]
    m2 = re.search(r"\bReferences\b|\bAcknowledg\w+\b", tail, re.I)
    if m2:
        tail = tail[: m2.start()]
    return tail.strip()


def _read_pages_text(pdf_path: Path, start: int, end: int) -> str:
    doc = fitz.open(pdf_path)
    try:
        parts = []
        for i in range(start, min(end, doc.page_count)):
            parts.append(doc.load_page(i).get_text("text"))
        return "\n".join(parts)
    finally:
        doc.close()


def _read_all_text(pdf_path: Path) -> str:
    doc = fitz.open(pdf_path)
    try:
        return "\n".join(doc.load_page(i).get_text("text") for i in range(doc.page_count))
    finally:
        doc.close()


def _add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = "Heading 1"


def _add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{key}: ")
    r1.bold = True
    p.add_run(value)


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


TRANSLATIONS: dict[str, dict[str, str]] = {
    "A_Review_of_Cyber-Ranges_and_Test-Beds_Current_and.pdf": {
        "core_zh": {
            "研究背景": "攻击自动化加速，IT/OT 边界收敛，组织需要更强的态势感知与训练来识别/应对新型威胁。",
            "研究问题": "现有网络靶场（CR）与测试床（TB）平台在类型、技术、威胁场景与训练能力上如何划分与对比？未来将如何演进？",
            "研究方法": "系统性综述与对比分析；按平台类型/技术/场景/应用/训练范围分段归纳，并构建 taxonomy。",
            "关键发现": "CR 与 TB 在应用领域的区分正在减弱，呈现收敛趋势；taxonomy 可用于理解与预测演化方向。",
            "实践启示": "训练宜在非生产环境开展，并提供贴近真实的威胁信息与对抗过程，支撑决策与防护能力提升。",
        },
    },
    "Cyber Operations RangE (CORE)_ Containerized Gaming Platform for.pdf": {
        "core_zh": {
            "研究背景": "CTF/攻防演练有助于教学与训练，但传统平台依赖复杂虚拟化，运维重且难以支撑实时对抗与计分。",
            "研究问题": "如何构建一种更易运维、可扩展且支持实时网络作战竞赛与独立计分的平台？",
            "研究方法": "设计并实现 CORE：将参赛环境容器化，按侦察/攻击/防御/取证组织任务流程。",
            "关键发现": "容器化显著降低管理负担，可用于不同规模竞赛，并更贴近实时对抗需求。",
            "实践启示": "可用于课堂体验式教学、组织内训练/评估与安全意识推广，也可服务定制化威胁分析与场景测试。",
        },
    },
    "Experimental Analysis of Security Attacks for Docker Container Communications.pdf": {
        "core_zh": {
            "研究背景": "Docker 容器广泛用于云与应用部署，容器间网络通信便利但也引入新的攻击面。",
            "研究问题": "容器间通信场景下，常见攻击（ARP 欺骗、DDoS、权限提升）如何发生，影响有多大？",
            "研究方法": "在容器网络环境中实验复现多类攻击，并从流量、CPU 与反向 shell 等维度观测效果。",
            "关键发现": "容器间通信确实可被上述攻击利用，影响体现在资源消耗与恶意控制链路建立等方面。",
            "实践启示": "需要面向容器间通信的检测与防护能力建设，并进一步评估安全机制在此类攻击下的有效性与局限。",
        },
    },
    "The Docker Security Playground A hands-on approach to the study of network security.pdf": {
        "core_zh": {
            "研究背景": "网络安全学习需要可重复、可组合的实验环境，传统搭建成本高且不易复用。",
            "研究问题": "如何用容器/微服务的方式快速构建复杂网络实验室，并支持教学与扩展开发？",
            "研究方法": "设计并实现 DSP：基于微服务的网络实验室架构，提供公共实验库与面向开发者的 API。",
            "关键发现": "该架构能更顺滑地创建/管理虚拟网络实验，便于学生动手实践与持续扩展新场景。",
            "实践启示": "后续可通过更高级的 compose 特性与实验室挂起/恢复机制增强生命周期管理，并持续丰富公开攻击场景库。",
        },
    },
}


def _build_docx(
    *,
    pdf_path: Path,
    out_path: Path,
    abstract_en: str | None,
    conclusion_en: str | None,
    abstract_zh: str | None,
    conclusion_zh: str | None,
    highlights_zh: str | None,
) -> None:
    doc = Document()
    _set_default_font(doc)

    _add_heading(doc, "英文文献中文翻译（摘译）")
    _add_kv(doc, "原文件", pdf_path.name)
    _add_kv(doc, "说明", "本文件为“核心要素摘要式翻译”（非全文逐段翻译）。")

    doc.add_paragraph()

    _add_heading(doc, "核心要素（中文摘要）")
    if highlights_zh:
        doc.add_paragraph(highlights_zh)
    else:
        doc.add_paragraph("（暂缺）")

    _add_heading(doc, "摘要（原文，自动提取）")
    doc.add_paragraph(abstract_en or "（未能自动提取摘要）")

    _add_heading(doc, "结论（原文，自动提取）")
    doc.add_paragraph(conclusion_en or "（未能自动提取结论）")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    base = Path("ctf/docs/文献")
    pdfs = [
        base / "A_Review_of_Cyber-Ranges_and_Test-Beds_Current_and.pdf",
        base / "Cyber Operations RangE (CORE)_ Containerized Gaming Platform for.pdf",
        base / "Experimental Analysis of Security Attacks for Docker Container Communications.pdf",
        base / "The Docker Security Playground A hands-on approach to the study of network security.pdf",
    ]

    for pdf in pdfs:
        first_pages = _read_pages_text(pdf, 0, 6)
        abstract_en = _extract_abstract(first_pages)
        full_text = _read_all_text(pdf)
        conclusion_en = _extract_conclusion(full_text)

        t = TRANSLATIONS.get(pdf.name, {})
        core_zh = t.get("core_zh") or {}
        core_lines = []
        for k in ["研究背景", "研究问题", "研究方法", "关键发现", "实践启示"]:
            v = core_zh.get(k)
            if v:
                core_lines.append(f"{k}：{v}")
        highlights_zh = "\n".join(core_lines) if core_lines else None

        out_name = pdf.stem + "_中文翻译.docx"
        out_path = pdf.with_name(out_name)
        _build_docx(
            pdf_path=pdf,
            out_path=out_path,
            abstract_en=abstract_en,
            conclusion_en=conclusion_en,
            highlights_zh=highlights_zh,
            abstract_zh=None,
            conclusion_zh=None,
        )
        print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
